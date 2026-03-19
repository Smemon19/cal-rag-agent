from __future__ import annotations

import os
import re
import io
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List

import google.auth
import streamlit as st
from vertexai import init as vertexai_init
from vertexai.generative_models import GenerativeModel

try:
    import fitz
except Exception:  # pragma: no cover - handled at runtime
    fitz = None  # type: ignore

try:
    import pdfplumber
except Exception:  # pragma: no cover - handled at runtime
    pdfplumber = None  # type: ignore

try:
    from docx import Document
except Exception:  # pragma: no cover - handled at runtime
    Document = None  # type: ignore


BUCKET_NAME = os.getenv("POLICY_PRO_BUCKET", "badgers-policy-docs")
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}
DEFAULT_TOP_K = min(max(int(os.getenv("POLICY_PRO_TOP_K", "5")), 1), 8)
DEFAULT_MAX_CONTEXT_CHARS = max(int(os.getenv("POLICY_PRO_MAX_CONTEXT_CHARS", "20000")), 2000)
DEFAULT_MIN_CHUNK_WORDS = max(int(os.getenv("POLICY_PRO_MIN_CHUNK_WORDS", "200")), 50)
DEFAULT_MAX_CHUNK_WORDS = max(int(os.getenv("POLICY_PRO_MAX_CHUNK_WORDS", "400")), DEFAULT_MIN_CHUNK_WORDS)
MIN_KEYWORD_CHARS = 2

STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "do", "for", "from", "how", "i", "in", "is", "it",
    "me", "my", "of", "on", "or", "our", "should", "that", "the", "their", "there", "they", "this", "to",
    "us", "was", "we", "what", "when", "where", "which", "who", "why", "with", "you", "your",
}


def _clean_extracted_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\b(?:[A-Za-z]\s+){2,}[A-Za-z]\b", lambda m: re.sub(r"\s+", "", m.group(0)), text)
    text = re.sub(r"\b(?:\d\s+){2,}\d\b", lambda m: re.sub(r"\s+", "", m.group(0)), text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_pdf_bytes(data: bytes, filename: str) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is required to read PDF files.")
    parts: List[str] = []

    # Primary extraction: PyMuPDF for full page text + annotations.
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        for page in doc:
            page_text = page.get_text("text") or ""
            if page_text.strip():
                parts.append(page_text)

            annot_texts: List[str] = []
            annots = page.annots()
            if annots:
                for annot in annots:
                    content = annot.info.get("content", "") if annot.info else ""
                    if content:
                        annot_texts.append(content)
            if annot_texts:
                parts.append("\n".join(annot_texts))
    finally:
        doc.close()

    # Fallback extraction for table-like text using pdfplumber.
    if pdfplumber is not None:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                for table in tables:
                    for row in table or []:
                        row_text = " | ".join((cell or "").strip() for cell in row)
                        if row_text.strip():
                            parts.append(row_text)

    return _clean_extracted_text("\n\n".join(parts))


def _read_docx_bytes(data: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required to read DOCX files.")
    doc = Document(io.BytesIO(data))
    out: List[str] = []

    for p in doc.paragraphs:
        if (p.text or "").strip():
            out.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
            if row_text.strip():
                out.append(row_text)

    for section in doc.sections:
        for p in section.header.paragraphs:
            if (p.text or "").strip():
                out.append(p.text)
        for p in section.footer.paragraphs:
            if (p.text or "").strip():
                out.append(p.text)

    return _clean_extracted_text("\n\n".join(out))


def _read_txt_bytes(data: bytes) -> str:
    return _clean_extracted_text(data.decode("utf-8", errors="ignore"))


def _extract_text_from_blob(name: str, data: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        return _read_pdf_bytes(data, name)
    if suffix == ".docx":
        return _read_docx_bytes(data)
    if suffix == ".txt":
        return _read_txt_bytes(data)
    return ""


def _words(text: str) -> List[str]:
    return re.findall(r"[a-z0-9']+", (text or "").lower())


def _extract_keywords(question: str) -> List[str]:
    seen: set[str] = set()
    out: List[str] = []
    for token in _words(question):
        if token in STOP_WORDS or len(token) < MIN_KEYWORD_CHARS:
            continue
        if token not in seen:
            seen.add(token)
            out.append(token)
    return out


def _split_long_paragraph(paragraph: str, max_words: int) -> List[str]:
    raw_words = paragraph.split()
    if not raw_words:
        return []
    chunks: List[str] = []
    idx = 0
    while idx < len(raw_words):
        chunks.append(" ".join(raw_words[idx : idx + max_words]).strip())
        idx += max_words
    return [c for c in chunks if c]


def _chunk_text_by_paragraphs(text: str, min_words: int = DEFAULT_MIN_CHUNK_WORDS, max_words: int = DEFAULT_MAX_CHUNK_WORDS) -> List[str]:
    if not text.strip():
        return []
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks: List[str] = []
    current_parts: List[str] = []
    current_words = 0

    def flush_current() -> None:
        nonlocal current_parts, current_words
        if current_parts:
            chunk_text = "\n\n".join(current_parts).strip()
            if chunk_text:
                chunks.append(chunk_text)
        current_parts = []
        current_words = 0

    for paragraph in paragraphs:
        para_word_count = len(paragraph.split())
        para_parts = [paragraph]
        if para_word_count > max_words:
            para_parts = _split_long_paragraph(paragraph, max_words=max_words)

        for para_part in para_parts:
            part_words = len(para_part.split())
            if part_words == 0:
                continue

            if current_words and (current_words + part_words > max_words):
                flush_current()

            current_parts.append(para_part)
            current_words += part_words

            if current_words >= min_words:
                flush_current()

    flush_current()
    return chunks


def _score_chunk(chunk: Dict[str, Any], keywords: List[str], lowered_question: str) -> tuple[int, int, int]:
    if not keywords:
        return (0, 0, 0)
    token_counts: Counter[str] = chunk["token_counts"]
    lowered = chunk["text_lc"]
    occurrences = sum(token_counts.get(k, 0) for k in keywords)
    unique_coverage = sum(1 for k in keywords if token_counts.get(k, 0) > 0)
    phrase_bonus = 1 if lowered_question and lowered_question in lowered else 0
    return (occurrences, unique_coverage, phrase_bonus)


def _build_context(top_chunks: List[Dict[str, Any]]) -> str:
    parts: List[str] = ["Policy Document Context"]
    for idx, chunk in enumerate(top_chunks, start=1):
        parts.append(
            f"[S{idx}] Document: {chunk['doc']}\n"
            f"Text:\n{chunk['text']}"
        )
    return "\n\n".join(parts)


@st.cache_resource(show_spinner=False)
def _cached_policy_pro_context() -> Dict[str, Any]:
    try:
        from google.cloud import storage
    except Exception as exc:
        raise RuntimeError(
            "google-cloud-storage is required for Policy Badger Pro. "
            "Install with: pip install google-cloud-storage"
        ) from exc

    client = storage.Client()
    bucket = client.bucket(BUCKET_NAME)
    blobs = sorted(bucket.list_blobs(), key=lambda b: b.name)
    documents: List[Dict[str, str]] = []
    chunks: List[Dict[str, Any]] = []
    chunk_id = 0

    print(f"[policy_pro] bucket=gs://{BUCKET_NAME} files_found={len(blobs)}")
    for blob in blobs:
        if blob.name.endswith("/"):
            continue
        suffix = Path(blob.name).suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue

        data = blob.download_as_bytes()
        text = _extract_text_from_blob(blob.name, data)
        char_count = len(text)
        print(f"[policy_pro] extracted file={blob.name} chars={char_count}")
        if char_count < 50:
            print(f"[policy_pro][warn] low extracted chars (<50) for file={blob.name}")

        if not text:
            continue

        documents.append({"filename": blob.name, "text": text})
        for piece in _chunk_text_by_paragraphs(text):
            tokens = _words(piece)
            if len(tokens) < 20:
                continue
            chunk_id += 1
            chunks.append(
                {
                    "chunk_id": f"chunk_{chunk_id}",
                    "doc": blob.name,
                    "text": piece,
                    "text_lc": piece.lower(),
                    "token_counts": Counter(tokens),
                    "char_count": len(piece),
                    "word_count": len(tokens),
                }
            )

    total_chars = sum(len(d["text"]) for d in documents)
    stats = {
        "documents_loaded": len(documents),
        "chunks_indexed": len(chunks),
        "total_chars": total_chars,
    }
    print(
        f"[policy_pro] documents={stats['documents_loaded']} "
        f"chunks={stats['chunks_indexed']} total_chars={stats['total_chars']}"
    )
    return {"documents": documents, "chunks": chunks, "stats": stats}


def load_policy_pro_context() -> Dict[str, Any]:
    return _cached_policy_pro_context()


def _select_relevant_chunks(
    question: str,
    context: Dict[str, Any],
    top_k: int = DEFAULT_TOP_K,
    max_context_chars: int = DEFAULT_MAX_CONTEXT_CHARS,
) -> Dict[str, Any]:
    top_k = min(max(top_k, 1), 8)
    max_context_chars = max(max_context_chars, 1000)
    chunks = context.get("chunks", [])
    keywords = _extract_keywords(question)
    lowered_question = question.strip().lower()
    ranked: List[tuple[tuple[int, int, int], Dict[str, Any]]] = []
    for chunk in chunks:
        score_tuple = _score_chunk(chunk, keywords, lowered_question)
        if score_tuple[0] <= 0 and score_tuple[2] <= 0:
            continue
        ranked.append((score_tuple, chunk))

    ranked.sort(
        key=lambda item: (
            item[0][0],  # keyword occurrences
            item[0][1],  # unique coverage
            item[0][2],  # phrase bonus
            -item[1]["char_count"],  # slightly prefer denser chunks
        ),
        reverse=True,
    )

    selected: List[Dict[str, Any]] = []
    running_chars = 0
    for _, chunk in ranked:
        if len(selected) >= top_k:
            break
        projected = running_chars + int(chunk["char_count"])
        if selected and projected > max_context_chars:
            break
        selected.append(chunk)
        running_chars = projected

    return {
        "keywords": keywords,
        "chunks": selected,
        "context_chars": running_chars,
    }


def ask_policy_pro(question: str, context: Dict[str, Any]) -> Dict[str, Any]:
    retrieval = _select_relevant_chunks(question, context)
    selected_chunks: List[Dict[str, Any]] = retrieval["chunks"]
    if not selected_chunks:
        return {
            "answer_text": (
                "I could not find that information in the retrieved policy excerpts. "
                "Please try rephrasing your question with more specific policy terms."
            ),
            "sources": [],
            "debug": {
                **context.get("stats", {}),
                "context_chunks_used": 0,
                "context_chars_used": 0,
                "keywords": retrieval["keywords"],
            },
        }

    creds, detected_project = google.auth.default()
    project_id = (
        os.getenv("VERTEX_PROJECT_ID")
        or os.getenv("GOOGLE_CLOUD_PROJECT")
        or detected_project
        or "badgers-487618"
    )
    location = os.getenv("VERTEX_LOCATION", "us-central1")
    model_name = os.getenv("CAL_MODEL_NAME", "gemini-2.0-flash")
    vertexai_init(project=project_id, location=location, credentials=creds)
    model = GenerativeModel(model_name)

    focused_context = _build_context(selected_chunks)
    prompt = (
        "You are an internal Raymond policy assistant. Use ONLY the provided policy document excerpts. "
        "Do not use outside knowledge. If the answer cannot be found in the excerpts, say "
        "\"The information was not found in the provided policy excerpts.\" "
        "Keep the answer concise and factual.\n\n"
        f"{focused_context}\n\n"
        "USER QUESTION:\n"
        f"{question}"
    )
    response = model.generate_content(prompt)
    answer = (response.text or "").strip()
    sources = sorted({chunk["doc"] for chunk in selected_chunks})
    return {
        "answer_text": answer,
        "sources": sources,
        "debug": {
            **context.get("stats", {}),
            "context_chunks_used": len(selected_chunks),
            "context_chars_used": retrieval["context_chars"],
            "keywords": retrieval["keywords"],
        },
    }
